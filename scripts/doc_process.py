"""Document preprocessing pipeline for medical RAG corpora.

This module ingests multi-format medical knowledge assets (PDF, Office, markdown,
HTML, text) and normalises them into clean text, tables and structured metadata
for use in retrieval-augmented generation (RAG) workflows. The implementation is
opinionated about medical-domain quirks such as reference lists, scanned PDFs,
multilingual material (Chinese/English) and traceability requirements.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import re
import unicodedata
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import textract
except ImportError:
    textract = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    import markdown as md
except ImportError:
    md = None

try:
    from langdetect import detect as lang_detect
except ImportError:
    lang_detect = None


DEFAULT_ENCODING_CANDIDATES: Tuple[str, ...] = (
    "utf-8",
    "utf-16",
    "utf-16-le",
    "utf-16-be",
    "gbk",
    "gb2312",
    "big5",
    "shift_jis",
    "latin-1",
)

REFERENCE_HEADERS: Tuple[str, ...] = (
    "references",
    "reference list",
    "bibliography",
    "citations",
    "参考文献",
    "參考文獻",
    "引用文献",
    "参考资料",
)

SUPPORTED_EXTENSIONS: Tuple[str, ...] = (
    ".pdf",
    ".docx",
    ".doc",
    ".md",
    ".markdown",
    ".txt",
    ".html",
    ".htm",
)

PDF_DATE_PATTERN = re.compile(
    r"D:(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})([+\-Zz].*)?"
)

DEFAULT_TEXT_OVERLAP = 300

logger = logging.getLogger("document_preprocessor")

class DocumentProcessingError(Exception):
    """Raised when a document cannot be processed."""


@dataclass
class TableArtifact:
    """Captures a table extracted from a document."""

    name: str
    data: List[List[str]]
    page: Optional[int] = None
    caption: Optional[str] = None

    def to_summary(self) -> Dict[str, Any]:
        """Return high-level table statistics for metadata serialization."""

        cols = max((len(row) for row in self.data), default=0)
        return {
            "name": self.name,
            "rows": len(self.data),
            "cols": cols,
            "page": self.page,
            "caption": self.caption,
        }


@dataclass
class DocumentMetadata:
    """Structured metadata payload recorded per source document."""

    document_id: str
    source_path: str
    source_name: str
    doc_type: str
    language: Optional[str] = None
    title: Optional[str] = None
    authors: List[str] = field(default_factory=list)
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    references: List[str] = field(default_factory=list)
    raw_metadata: Dict[str, str] = field(default_factory=dict)
    extra: Dict[str, str] = field(default_factory=dict)

    def to_serializable(self) -> Dict[str, Any]:
        """Transform metadata into a JSON-serialisable dictionary."""

        return {
            "document_id": self.document_id,
            "source_path": self.source_path,
            "source_name": self.source_name,
            "doc_type": self.doc_type,
            "language": self.language,
            "title": self.title,
            "authors": self.authors,
            "created_at": self.created_at,
            "modified_at": self.modified_at,
            "references": self.references,
            "raw_metadata": self.raw_metadata,
            "extra": self.extra,
        }


@dataclass
class DocumentPayload:
    """Container holding extracted text, metadata and optional tables."""

    content: str
    metadata: DocumentMetadata
    tables: List[TableArtifact] = field(default_factory=list)
    relative_dir: Optional[Path] = None


@dataclass
class PreprocessConfig:
    """Runtime switches controlling document preprocessing behaviour."""

    input_dir: Path
    output_dir: Path
    languages_hint: Optional[Sequence[str]] = None
    skip_ocr: bool = False
    maintain_relative_structure: bool = True
    overwrite: bool = False
    fail_fast: bool = False
    encoding_candidates: Sequence[str] = DEFAULT_ENCODING_CANDIDATES
    allowed_extensions: Sequence[str] = SUPPORTED_EXTENSIONS
    text_subdir: str = "clean_text"
    metadata_subdir: str = "metadata"
    table_subdir: str = "tables"
    max_text_chars: Optional[int] = None
    text_overlap: int = DEFAULT_TEXT_OVERLAP

    def __post_init__(self) -> None:
        """Normalise user-provided switches and prepare target directories."""

        self.input_dir = Path(self.input_dir)
        self.output_dir = Path(self.output_dir)
        self.text_dir = self.output_dir / self.text_subdir
        self.metadata_dir = self.output_dir / self.metadata_subdir
        self.table_dir = self.output_dir / self.table_subdir
        self.allowed_extensions = tuple(ext.lower() for ext in self.allowed_extensions)
        if self.languages_hint:
            self.languages_hint = [lang.strip() for lang in self.languages_hint if lang]
        else:
            self.languages_hint = None
        if isinstance(self.max_text_chars, str):
            self.max_text_chars = int(self.max_text_chars)
        if isinstance(self.text_overlap, str):
            self.text_overlap = int(self.text_overlap)
        if self.max_text_chars is not None and self.max_text_chars <= 0:
            self.max_text_chars = None
        if self.text_overlap < 0:
            self.text_overlap = 0
        for directory in (
            self.output_dir,
            self.text_dir,
            self.metadata_dir,
            self.table_dir,
        ):
            directory.mkdir(parents=True, exist_ok=True)

def slugify(value: str, allow_unicode: bool = True) -> str:
    """Generate a filesystem-safe slug from a string."""

    value = str(value)
    if allow_unicode:
        value = unicodedata.normalize("NFKC", value)
    else:
        value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^\w\s-]", "", value, flags=re.UNICODE)
    value = re.sub(r"[-\s]+", "-", value).strip("-_")
    return value


def normalize_text(text: str) -> str:
    """Collapse Windows newlines and excessive whitespace for downstream use."""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_metadata_dict(raw: Dict[str, Any]) -> Dict[str, str]:
    """Convert metadata dictionaries into plain-string key/value mappings."""

    normalized: Dict[str, str] = {}
    for key, value in raw.items():
        if value is None:
            continue
        normalized_key = str(key)
        if normalized_key.startswith("/"):
            normalized_key = normalized_key[1:]
        if isinstance(value, (list, tuple, set)):
            flattened = [str(item) for item in value if item is not None]
            if not flattened:
                continue
            normalized[normalized_key] = ", ".join(flattened)
        elif isinstance(value, (bytes, bytearray)):
            normalized[normalized_key] = value.decode("utf-8", errors="ignore")
        else:
            normalized[normalized_key] = str(value)
    return normalized


def compute_sha256(path: Path, chunk_size: int = 65536) -> str:
    """Stream file content and return the SHA-256 checksum."""

    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(chunk_size), b""):
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def detect_language(text: str, hints: Optional[Sequence[str]] = None) -> Optional[str]:
    """Use langdetect on a text snippet, falling back to configured hints."""

    sample = text.strip()
    if not sample:
        return hints[0] if hints else None
    sample = sample[:2000]
    if lang_detect is None:
        return hints[0] if hints else None
    try:
        return lang_detect(sample)
    except Exception:
        return hints[0] if hints else None


def extract_references(full_text: str) -> List[str]:
    """Locate reference sections using multilingual headers and return entries."""

    references: List[str] = []
    lines = [line.strip() for line in full_text.splitlines()]
    start_index: Optional[int] = None
    for idx, line in enumerate(lines):
        if not line:
            continue
        normalized = re.sub(r"[：:]+$", "", line.lower())
        if start_index is None and any(
            normalized.startswith(header) for header in REFERENCE_HEADERS
        ):
            start_index = idx + 1
            continue
        if start_index is not None and idx >= start_index:
            if re.match(r"^[A-Z][\w\s\-:]{0,60}$", line) and len(references) > 5:
                break
            references.append(line)
    cleaned = [ref for ref in references if ref]
    return cleaned[:200]


def read_text_file(path: Path, encodings: Sequence[str]) -> str:
    """Attempt multiple encodings when loading a plain-text file."""

    data = path.read_bytes()
    for encoding in encodings:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def parse_pdf_timestamp(value: str) -> Optional[str]:
    """Convert PDF timestamp strings into ISO 8601, when possible."""

    match = PDF_DATE_PATTERN.match(value)
    if not match:
        return None
    try:
        year, month, day, hour, minute, second, zone = match.groups()
        dt = datetime(
            int(year),
            int(month),
            int(day),
            int(hour),
            int(minute),
            int(second),
        )
        if zone and zone.upper() != "Z":
            zone = zone.replace("'", "")
            return f"{dt.isoformat()}{zone}"
        return f"{dt.isoformat()}Z"
    except ValueError:
        return None


def chunk_text(content: str, max_chars: Optional[int], overlap: int) -> List[str]:
    """Split long text into overlapping chunks bounded by `max_chars` characters."""

    if not content:
        return []
    if max_chars is None or max_chars <= 0 or len(content) <= max_chars:
        return [content]
    overlap = max(0, overlap)
    if overlap >= max_chars:
        overlap = max_chars // 5
    chunks: List[str] = []
    start = 0
    length = len(content)
    while start < length:
        end = min(start + max_chars, length)
        if end < length:
            candidate = content.rfind("\n\n", start + max_chars // 2, end)
            if candidate == -1:
                candidate = content.rfind("\n", start + max_chars // 2, end)
            if candidate != -1 and candidate > start:
                end = candidate
        segment = content[start:end].strip()
        if not segment:
            segment = content[start:end]
        if segment:
            chunks.append(segment)
        if end >= length:
            break
        start = max(end - overlap, start + 1)
    return chunks

class DocumentPreprocessor:
    """Core coordinator orchestrating decoding, metadata capture and persistence."""

    def __init__(self, config: PreprocessConfig) -> None:
        self.config = config
        self.logger = logging.getLogger("document_preprocessor")
        self.handlers = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_doc,
            ".md": self._process_markdown,
            ".markdown": self._process_markdown,
            ".txt": self._process_text,
            ".html": self._process_html,
            ".htm": self._process_html,
        }
        self.ocr_language_hint = self._resolve_ocr_language_hint()

    def run(self) -> Dict[str, List[str]]:
        """Iterate over all files in the input tree and process those supported."""

        processed: List[str] = []
        failures: List[str] = []
        for path in self._iter_documents():
            self.logger.info("Processing %s", path)
            try:
                payload = self._process_path(path)
                self._persist_payload(path, payload)
            except DocumentProcessingError as exc:
                self.logger.error("Failed to process %s: %s", path, exc)
                failures.append(str(path))
                if self.config.fail_fast:
                    raise
            except Exception as exc:
                self.logger.exception("Unexpected error processing %s: %s", path, exc)
                failures.append(str(path))
                if self.config.fail_fast:
                    raise
            else:
                processed.append(str(path))
        self.logger.info(
            "Finished preprocessing: %s processed, %s failed",
            len(processed),
            len(failures),
        )
        return {"processed": processed, "failed": failures}

    def _iter_documents(self) -> Iterable[Path]:
        """Yield candidate files underneath the input directory."""

        for path in sorted(self.config.input_dir.rglob("*")):
            if not path.is_file():
                continue
            if path.suffix.lower() in self.config.allowed_extensions:
                yield path

    def _process_path(self, path: Path) -> DocumentPayload:
        """Dispatch to the correct handler and enrich the metadata payload."""

        handler = self.handlers.get(path.suffix.lower())
        if handler is None:
            raise DocumentProcessingError(f"Unsupported extension: {path.suffix}")
        document_id = self._build_document_id(path)
        payload = handler(path, document_id)
        payload.content = normalize_text(payload.content)
        if not payload.metadata.references:
            payload.metadata.references = extract_references(payload.content)
        if not payload.metadata.language:
            payload.metadata.language = detect_language(
                payload.content, self.config.languages_hint
            )
        payload.metadata.extra["tables_detected"] = str(len(payload.tables))
        payload.metadata.extra["content_char_count"] = str(len(payload.content))
        payload.metadata.extra["source_filesize_bytes"] = str(path.stat().st_size)
        payload.metadata.extra["source_checksum_sha256"] = compute_sha256(path)
        if self.config.max_text_chars:
            payload.metadata.extra["text_chunk_strategy"] = (
                f"max_chars:{self.config.max_text_chars} overlap:{self.config.text_overlap}"
            )
        rel_dir = self._resolve_relative_dir(path)
        payload.relative_dir = rel_dir
        if rel_dir is not None:
            relative_source = str(rel_dir / path.name) if rel_dir != Path() else path.name
            payload.metadata.extra["relative_source_path"] = relative_source
        payload.metadata.extra["processed_at"] = (
            datetime.utcnow().isoformat(timespec="seconds") + "Z"
        )
        return payload

    def _resolve_relative_dir(self, path: Path) -> Optional[Path]:
        """Respect nested input layout if configured to maintain structure."""

        if not self.config.maintain_relative_structure:
            return None
        try:
            relative = path.parent.relative_to(self.config.input_dir)
        except ValueError:
            return None
        return relative

    def _persist_payload(self, source_path: Path, payload: DocumentPayload) -> None:
        """Create text/tables/metadata outputs in the configured directories."""

        rel_dir = payload.relative_dir if payload.relative_dir is not None else Path()
        text_dir = self.config.text_dir / rel_dir
        metadata_dir = self.config.metadata_dir / rel_dir
        table_dir = self.config.table_dir / rel_dir
        for directory in (text_dir, metadata_dir, table_dir):
            directory.mkdir(parents=True, exist_ok=True)
        text_paths, chunk_records = self._write_text_outputs(payload, text_dir)
        table_entries: List[Dict[str, Any]] = []
        for index, table in enumerate(payload.tables, start=1):
            table_name = table.name or f"{payload.metadata.document_id}_table_{index}.csv"
            table_path = table_dir / table_name
            with table_path.open("w", newline="", encoding="utf-8") as handle:
                writer = csv.writer(handle)
                for row in table.data:
                    writer.writerow([cell if cell is not None else "" for cell in row])
            table_summary = table.to_summary()
            table_summary["path"] = self._to_output_relative(table_path)
            table_entries.append(table_summary)
        metadata_payload = payload.metadata.to_serializable()
        metadata_payload["text_path"] = (
            self._to_output_relative(text_paths[0]) if text_paths else None
        )
        metadata_payload["text_paths"] = [self._to_output_relative(path) for path in text_paths]
        metadata_payload["text_chunk_count"] = len(text_paths)
        metadata_payload["text_chunks"] = [
            {"seq": record["seq"], "path": self._to_output_relative(record["path"]), "char_count": record["char_count"]}
            for record in chunk_records
        ]
        metadata_payload["tables"] = table_entries
        metadata_payload["generated_at"] = (
            datetime.utcnow().isoformat(timespec="seconds") + "Z"
        )
        metadata_payload["source_extension"] = source_path.suffix.lower()
        metadata_payload["source_filename"] = source_path.name
        metadata_path = metadata_dir / f"{payload.metadata.document_id}.json"
        if metadata_path.exists() and not self.config.overwrite:
            raise DocumentProcessingError(
                f"Target metadata file already exists: {metadata_path}"
            )
        metadata_path.write_text(
            json.dumps(metadata_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _write_text_outputs(
        self, payload: DocumentPayload, text_dir: Path
    ) -> Tuple[List[Path], List[Dict[str, Any]]]:
        """Write cleaned text to disk, applying chunking when configured."""

        chunks = chunk_text(
            payload.content,
            self.config.max_text_chars,
            self.config.text_overlap,
        )
        if not chunks:
            chunks = [""]
        text_paths: List[Path] = []
        chunk_records: List[Dict[str, Any]] = []
        multiple = len(chunks) > 1
        for idx, chunk in enumerate(chunks, start=1):
            if multiple:
                file_name = f"{payload.metadata.document_id}_part{idx:02d}.txt"
            else:
                file_name = f"{payload.metadata.document_id}.txt"
            text_path = text_dir / file_name
            if text_path.exists() and not self.config.overwrite:
                raise DocumentProcessingError(
                    f"Target text file already exists: {text_path}"
                )
            text_path.write_text(chunk, encoding="utf-8")
            text_paths.append(text_path)
            chunk_records.append(
                {"seq": idx, "path": text_path, "char_count": len(chunk)}
            )
        payload.metadata.extra["text_chunk_count"] = str(len(chunks))
        return text_paths, chunk_records

    def _to_output_relative(self, path: Path) -> str:
        """Return a path relative to the output root when possible."""

        try:
            return str(path.relative_to(self.config.output_dir))
        except ValueError:
            return str(path)

    def _build_document_id(self, path: Path) -> str:
        """Generate a stable yet collision-resistant document identifier."""

        slug = slugify(path.stem)
        suffix = uuid.uuid4().hex[:8]
        if slug:
            return f"{slug}-{suffix}"
        return suffix

    def _resolve_ocr_language_hint(self) -> str:
        """Translate language hints into Tesseract language codes."""

        if not self.config.languages_hint:
            return "eng"
        mapping = {
            "en": "eng",
            "eng": "eng",
            "english": "eng",
            "zh": "chi_sim",
            "zh-cn": "chi_sim",
            "zh-hans": "chi_sim",
            "zh_cn": "chi_sim",
            "zh-hant": "chi_tra",
            "zh-tw": "chi_tra",
            "zh_tw": "chi_tra",
            "chi": "chi_sim",
        }
        codes: List[str] = []
        for hint in self.config.languages_hint:
            normalized = hint.lower()
            code = mapping.get(normalized) or mapping.get(normalized.replace("-", "_"))
            if code is None and "-" in normalized:
                prefix = normalized.split("-", 1)[0]
                code = mapping.get(prefix, prefix)
            if code is None:
                code = normalized
            if code not in codes:
                codes.append(code)
        return "+".join(codes) if codes else "eng"

    def _process_pdf(self, path: Path, document_id: str) -> DocumentPayload:
        """Extract text, tables and metadata from PDF files with OCR fallback."""

        if pdfplumber is None and PdfReader is None:
            raise DocumentProcessingError(
                "Processing PDF files requires pdfplumber or pypdf. "
                "Please install at least one of them."
            )
        text_segments: List[str] = []
        tables: List[TableArtifact] = []
        raw_meta: Dict[str, str] = {}
        page_count = 0
        extraction_method = "unknown"
        if pdfplumber is not None:
            try:
                with pdfplumber.open(str(path)) as pdf:
                    page_count = len(pdf.pages)
                    raw_meta.update(normalize_metadata_dict(pdf.metadata or {}))
                    extraction_method = "pdfplumber"
                    for page_index, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text() or ""
                        if page_text:
                            text_segments.append(page_text)
                        try:
                            extracted_tables = page.extract_tables()
                        except Exception:
                            extracted_tables = []
                        for table_index, table_data in enumerate(
                            extracted_tables, start=1
                        ):
                            if not table_data:
                                continue
                            cleaned = [
                                [
                                    cell.strip() if isinstance(cell, str) else ""
                                    for cell in row
                                ]
                                for row in table_data
                            ]
                            table_name = (
                                f"{document_id}_p{page_index}_t{table_index}.csv"
                            )
                            tables.append(
                                TableArtifact(
                                    name=table_name,
                                    data=cleaned,
                                    page=page_index,
                                )
                            )
            except Exception as exc:
                self.logger.warning(
                    "pdfplumber failed for %s (%s), attempting PyPDF fallback",
                    path,
                    exc,
                )
                text_segments = []
                tables = []
                raw_meta = {}
                page_count = 0
                extraction_method = "unknown"
        if not text_segments and PdfReader is not None:
            try:
                reader = PdfReader(str(path))
            except Exception as exc:
                raise DocumentProcessingError(f"Failed to read PDF: {exc}") from exc
            pdf_meta = reader.metadata or {}
            raw_meta.update(normalize_metadata_dict(pdf_meta))
            extraction_method = "pypdf"
            page_count = len(reader.pages)
            for page_index, page in enumerate(reader.pages, start=1):
                try:
                    text_segments.append(page.extract_text() or "")
                except Exception:
                    text_segments.append("")
        text_content = "\n\n".join(segment for segment in text_segments if segment).strip()
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="pdf",
            raw_metadata=raw_meta,
        )
        if not text_content and not self.config.skip_ocr:
            ocr_text, ocr_pages = self._run_pdf_ocr(path)
            if ocr_text:
                self.logger.info("Applied OCR fallback to %s", path)
                text_content = ocr_text
                metadata.extra["ocr_applied"] = "true"
                metadata.extra["ocr_language_hint"] = self.ocr_language_hint
                if not page_count:
                    page_count = ocr_pages
                extraction_method = "ocr"
        if not text_content:
            self.logger.warning("No textual content extracted from %s", path)
        if not metadata.title:
            metadata.title = raw_meta.get("Title") or raw_meta.get("title")
        author = raw_meta.get("Author") or raw_meta.get("author")
        if author:
            metadata.authors = [author]
        created = raw_meta.get("CreationDate") or raw_meta.get("creationdate")
        if created:
            parsed = parse_pdf_timestamp(created)
            metadata.created_at = parsed or created
        modified = raw_meta.get("ModDate") or raw_meta.get("moddate")
        if modified:
            parsed = parse_pdf_timestamp(modified)
            metadata.modified_at = parsed or modified
        producer = raw_meta.get("Producer") or raw_meta.get("producer")
        if producer:
            metadata.extra["pdf_producer"] = producer
        metadata.extra["page_count"] = str(page_count)
        metadata.extra["extraction_method"] = extraction_method
        return DocumentPayload(content=text_content, metadata=metadata, tables=tables)

    def _run_pdf_ocr(self, path: Path) -> Tuple[str, int]:
        """Fallback OCR for scanned PDFs using pdf2image+pytesseract."""

        try:
            from pdf2image import convert_from_path
        except ImportError:
            self.logger.warning(
                "Skipping OCR for %s because pdf2image is not installed.", path
            )
            return "", 0
        try:
            import pytesseract
        except ImportError:
            self.logger.warning(
                "Skipping OCR for %s because pytesseract is not installed.", path
            )
            return "", 0
        try:
            images = convert_from_path(str(path))
        except Exception as exc:
            self.logger.warning("Failed to render PDF %s for OCR: %s", path, exc)
            return "", 0
        texts: List[str] = []
        for index, image in enumerate(images, start=1):
            try:
                text = pytesseract.image_to_string(
                    image, lang=self.ocr_language_hint
                )
            except Exception as exc:
                self.logger.warning(
                    "OCR failed on page %s of %s: %s", index, path, exc
                )
                text = ""
            texts.append(text)
        combined = "\n\n".join(segment.strip() for segment in texts if segment).strip()
        return combined, len(images)

    def _process_docx(self, path: Path, document_id: str) -> DocumentPayload:
        """Extract paragraphs and tables from DOCX files."""

        if Document is None:
            raise DocumentProcessingError(
                "Processing .docx files requires python-docx. Please install it."
            )
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise DocumentProcessingError(f"Failed to open DOCX: {exc}") from exc
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        text_content = "\n\n".join(paragraphs)
        tables: List[TableArtifact] = []
        for table_index, table in enumerate(doc.tables, start=1):
            table_data: List[List[str]] = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            if not table_data:
                continue
            table_name = f"{document_id}_t{table_index}.csv"
            tables.append(TableArtifact(name=table_name, data=table_data))
        props = doc.core_properties
        raw_meta = normalize_metadata_dict(
            {
                "author": props.author,
                "category": props.category,
                "comments": props.comments,
                "content_status": props.content_status,
                "created": props.created,
                "identifier": props.identifier,
                "keywords": props.keywords,
                "language": props.language,
                "last_modified_by": props.last_modified_by,
                "last_printed": props.last_printed,
                "modified": props.modified,
                "subject": props.subject,
                "title": props.title,
                "version": props.version,
            }
        )
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="docx",
            raw_metadata=raw_meta,
        )
        metadata.title = props.title or raw_meta.get("title")
        if props.author:
            metadata.authors = [props.author]
        if props.created:
            metadata.created_at = props.created.isoformat()
        if props.modified:
            metadata.modified_at = props.modified.isoformat()
        if props.keywords:
            metadata.extra["keywords"] = props.keywords
        metadata.language = props.language
        metadata.extra["extraction_method"] = "python-docx"
        return DocumentPayload(content=text_content, metadata=metadata, tables=tables)

    def _process_doc(self, path: Path, document_id: str) -> DocumentPayload:
        """Handle legacy .doc files using textract when available."""

        if textract is None:
            raise DocumentProcessingError(
                "Processing legacy .doc files requires textract. Please install it."
            )
        try:
            binary = textract.process(str(path))
        except Exception as exc:
            raise DocumentProcessingError(f"textract failed: {exc}") from exc
        text_content = binary.decode("utf-8", errors="ignore")
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="doc",
        )
        metadata.extra[
            "notice"
        ] = "Tables are not automatically extracted from legacy .doc files."
        metadata.extra["extraction_method"] = "textract"
        return DocumentPayload(content=text_content, metadata=metadata)

    def _process_markdown(self, path: Path, document_id: str) -> DocumentPayload:
        """Inline markdown files, preserving tables via HTML parsing."""

        text_content = read_text_file(path, self.config.encoding_candidates)
        tables: List[TableArtifact] = []
        if md is not None and BeautifulSoup is not None:
            try:
                html = md.markdown(text_content, extensions=["tables"])
                tables = self._extract_tables_from_html(html, document_id)
            except Exception as exc:
                self.logger.debug(
                    "Markdown table extraction failed for %s: %s", path, exc
                )
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="markdown",
        )
        metadata.extra["extraction_method"] = "markdown"
        return DocumentPayload(content=text_content, metadata=metadata, tables=tables)

    def _process_text(self, path: Path, document_id: str) -> DocumentPayload:
        """Load plain-text resources as-is with encoding fallback."""

        text_content = read_text_file(path, self.config.encoding_candidates)
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="text",
        )
        metadata.extra["extraction_method"] = "text"
        return DocumentPayload(content=text_content, metadata=metadata)

    def _process_html(self, path: Path, document_id: str) -> DocumentPayload:
        """Strip scripts/styles and capture meta tags from HTML resources."""

        html_content = read_text_file(path, self.config.encoding_candidates)
        if BeautifulSoup is None:
            metadata = DocumentMetadata(
                document_id=document_id,
                source_path=str(path),
                source_name=path.name,
                doc_type="html",
            )
            metadata.extra["extraction_method"] = "raw-html"
            return DocumentPayload(content=html_content, metadata=metadata)
        soup = BeautifulSoup(html_content, "html.parser")
        for unwanted in soup(["script", "style", "noscript"]):
            unwanted.decompose()
        text_content = soup.get_text("\n", strip=True)
        tables = self._extract_tables_from_soup(soup, document_id)
        meta_info: Dict[str, Any] = {}
        for meta in soup.find_all("meta"):
            key = meta.get("name") or meta.get("property")
            value = meta.get("content")
            if key and value:
                meta_info[key] = value
        raw_meta = normalize_metadata_dict(meta_info)
        metadata = DocumentMetadata(
            document_id=document_id,
            source_path=str(path),
            source_name=path.name,
            doc_type="html",
            raw_metadata=raw_meta,
        )
        if soup.title and soup.title.string:
            metadata.title = soup.title.string.strip()
        html_lang = soup.html.get("lang") if soup.html else None
        if html_lang:
            metadata.language = html_lang
        metadata.extra["extraction_method"] = "beautifulsoup"
        return DocumentPayload(content=text_content, metadata=metadata, tables=tables)

    def _extract_tables_from_html(
        self, html: str, document_id: str
    ) -> List[TableArtifact]:
        """Convert HTML tables rendered from markdown into CSV artifacts."""

        if BeautifulSoup is None:
            return []
        soup = BeautifulSoup(html, "html.parser")
        return self._extract_tables_from_soup(soup, document_id)

    def _extract_tables_from_soup(
        self, soup: Any, document_id: str
    ) -> List[TableArtifact]:
        """Walk HTML soup and extract tabular data with optional captions."""

        tables: List[TableArtifact] = []
        for index, table_tag in enumerate(soup.find_all("table"), start=1):
            rows: List[List[str]] = []
            for row_tag in table_tag.find_all("tr"):
                cells = row_tag.find_all(["th", "td"])
                if not cells:
                    continue
                rows.append([cell.get_text(strip=True) for cell in cells])
            if not rows:
                continue
            caption_tag = table_tag.find("caption")
            caption = caption_tag.get_text(strip=True) if caption_tag else None
            table_name = f"{document_id}_t{index}.csv"
            tables.append(
                TableArtifact(name=table_name, data=rows, caption=caption)
            )
        return tables

def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    """CLI argument definition for the preprocessing utility."""

    parser = argparse.ArgumentParser(
        description="Preprocess medical documents for RAG ingestion."
    )
    parser.add_argument(
        "--input-dir",
        required=True,
        help="Directory containing source documents.",
    )
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Directory to store cleaned text, tables, and metadata.",
    )
    parser.add_argument(
        "--language-hint",
        action="append",
        dest="languages_hint",
        default=None,
        help="Language hint for extraction and OCR (repeatable).",
    )
    parser.add_argument(
        "--skip-ocr",
        action="store_true",
        help="Skip OCR fallback for scanned PDFs.",
    )
    parser.add_argument(
        "--no-relative-structure",
        action="store_true",
        help="Do not preserve input subdirectory structure in outputs.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing output files if they exist.",
    )
    parser.add_argument(
        "--fail-fast",
        action="store_true",
        help="Stop processing when the first document fails.",
    )
    parser.add_argument(
        "--max-text-chars",
        type=int,
        default=None,
        help="Maximum characters per text chunk before splitting.",
    )
    parser.add_argument(
        "--text-overlap",
        type=int,
        default=DEFAULT_TEXT_OVERLAP,
        help="Character overlap between neighbouring text chunks.",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        help="Logging level (e.g., INFO, DEBUG).",
    )
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> None:
    """Entrypoint for CLI execution."""

    args = parse_args(argv)
    logging.basicConfig(
        level=getattr(logging, args.log_level.upper(), logging.INFO),
        format="%(asctime)s [%(levelname)s] %(message)s",
    )
    config = PreprocessConfig(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        languages_hint=args.languages_hint,
        skip_ocr=args.skip_ocr,
        maintain_relative_structure=not args.no_relative_structure,
        overwrite=args.overwrite,
        fail_fast=args.fail_fast,
        max_text_chars=args.max_text_chars,
        text_overlap=args.text_overlap,
    )
    processor = DocumentPreprocessor(config)
    summary = processor.run()
    logger.info(
        "Preprocessing summary:\n%s",
        json.dumps(summary, ensure_ascii=False, indent=2),
    )


if __name__ == "__main__":
    main()
